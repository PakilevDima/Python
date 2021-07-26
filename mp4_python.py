hiddenimports = ['speech_recognition', 'docx']


import speech_recognition
import docx
import tkinter


def record_volume():
    r = speech_recognition.Recognizer()
    with speech_recognition.Microphone(device_index=1) as source:
        print('Слушаю ...')
        try:
            audio = r.listen(source, timeout=4)
            query = r.recognize_google(audio, language='ru-RU')
            print(f'Вы сказали\n{query.lower()}')
        except (speech_recognition.WaitTimeoutError, speech_recognition.UnknownValueError):
            print('Извините, но запись не получилась:)). Говорите тётче')
        else:
            print('Запись получилась. Хотите записать??')
            string = input()
            if string == 'да':
                doc = docx.Document('звукозапись.docx')
                text = []
                for paragraph in doc.paragraphs:
                    text.append(paragraph.text)
                mydoc = docx.Document()
                mydoc.add_paragraph('\n'.join(text))
                mydoc.add_paragraph(query.lower())
                mydoc.save("звукозапись.docx")
                print('Файл успешно записан')
record_volume()
